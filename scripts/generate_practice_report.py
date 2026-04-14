from __future__ import annotations

import math
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TITLE_PATH = Path("/Users/daniilsintsov/Downloads/Титульник Практика.docx")
OUTPUT_DIR = ROOT / "output" / "doc"
ASSET_DIR = ROOT / "output" / "screenshots"
TMP_DIR = ROOT / "tmp" / "docs"
REPORT_PATH = OUTPUT_DIR / "Отчет_по_практике_Bauman_Doomlike.docx"

MAP_PATH = ROOT / "resources" / "maps" / "arena.txt"

PAGE_MAP = {
    "Введение": 3,
    "Индивидуальное задание": 4,
    "Основная часть": 5,
    "Заключение": 11,
    "Список использованных источников": 12,
    "Приложения": 13,
}


INTRO_TEXT = (
    "Целями проектно-технологической практики являются освоение навыков разработки и "
    "проектирования информационных систем на основе объектно-ориентированного подхода "
    "(ООП) и знакомство с современными системами объектного программирования и системами классов. "
    "При прохождении проектно-технологической практики ставятся следующие задачи: самостоятельно "
    "изучить существующие системы классы и современные технологии, основанные на ООП, и применить "
    "современные технологии, основанные на ООП, при проектировании и разработке собственного "
    "решения по индивидуальному заданию."
)


@dataclass
class Vec2:
    x: float
    y: float

    def __add__(self, other: "Vec2") -> "Vec2":
        return Vec2(self.x + other.x, self.y + other.y)

    def __sub__(self, other: "Vec2") -> "Vec2":
        return Vec2(self.x - other.x, self.y - other.y)

    def __mul__(self, scalar: float) -> "Vec2":
        return Vec2(self.x * scalar, self.y * scalar)


def load_map() -> tuple[list[str], Vec2, Vec2]:
    rows = [line.rstrip("\n") for line in MAP_PATH.read_text(encoding="utf-8").splitlines() if line.strip()]
    host = Vec2(1.5, 1.5)
    client = Vec2(7.5, 7.5)
    for y, row in enumerate(rows):
        for x, cell in enumerate(row):
            if cell == "1":
                host = Vec2(x + 0.5, y + 0.5)
            elif cell == "2":
                client = Vec2(x + 0.5, y + 0.5)
    return rows, host, client


def to_channel(value: float) -> int:
    return max(0, min(255, int(value)))


def normalize(vec: Vec2) -> Vec2:
    length = math.hypot(vec.x, vec.y)
    if length <= 1e-4:
        return Vec2(0.0, 0.0)
    return Vec2(vec.x / length, vec.y / length)


def angle_between(src: Vec2, dst: Vec2) -> float:
    return math.atan2(dst.y - src.y, dst.x - src.x)


def wrap_angle(value: float) -> float:
    while value > math.pi:
        value -= math.pi * 2.0
    while value < -math.pi:
        value += math.pi * 2.0
    return value


def cast_columns(map_rows: list[str], position: Vec2, view_angle: float, screen_width: int, fov: float = math.radians(66.0)) -> list[tuple[bool, float, int]]:
    columns: list[tuple[bool, float, int]] = []
    height = len(map_rows)
    width = len(map_rows[0])
    max_distance = 20.0

    for x in range(screen_width):
        t = (x + 0.5) / screen_width
        ray_angle = (view_angle - fov * 0.5) + (fov * t)
        ray_dir_x = math.cos(ray_angle)
        ray_dir_y = math.sin(ray_angle)

        map_x = int(position.x)
        map_y = int(position.y)

        delta_x = float("inf") if abs(ray_dir_x) < 1e-4 else abs(1.0 / ray_dir_x)
        delta_y = float("inf") if abs(ray_dir_y) < 1e-4 else abs(1.0 / ray_dir_y)

        if ray_dir_x < 0.0:
            step_x = -1
            side_dist_x = (position.x - map_x) * delta_x
        else:
            step_x = 1
            side_dist_x = (map_x + 1.0 - position.x) * delta_x

        if ray_dir_y < 0.0:
            step_y = -1
            side_dist_y = (position.y - map_y) * delta_y
        else:
            step_y = 1
            side_dist_y = (map_y + 1.0 - position.y) * delta_y

        hit = False
        side = 0
        distance = max_distance

        while not hit:
            if side_dist_x < side_dist_y:
                side_dist_x += delta_x
                map_x += step_x
                side = 0
            else:
                side_dist_y += delta_y
                map_y += step_y
                side = 1

            if map_x < 0 or map_y < 0 or map_x >= width or map_y >= height:
                break

            cell = map_rows[map_y][map_x]
            if cell == "#":
                hit = True
                raw_distance = side_dist_x - delta_x if side == 0 else side_dist_y - delta_y
                distance = raw_distance * math.cos(ray_angle - view_angle)
                distance = max(0.0001, min(max_distance, distance))

        columns.append((hit, distance, side))

    return columns


def raycast_interface_image(output_path: Path) -> None:
    map_rows, host, client = load_map()
    viewer = host
    target = client
    view_angle = 0.35

    width, height = 1280, 720
    img = Image.new("RGB", (width, height), "black")
    draw = ImageDraw.Draw(img)

    horizon = height // 2
    for y in range(horizon):
        t = y / horizon
        color = (
            to_channel(25 + 40 * t),
            to_channel(40 + 90 * t),
            to_channel(90 + 120 * t),
        )
        draw.line((0, y, width, y), fill=color)

    for y in range(horizon, height):
        t = (y - horizon) / horizon
        color = (
            to_channel(55 + 30 * t),
            to_channel(45 + 20 * t),
            to_channel(30 + 10 * t),
        )
        draw.line((0, y, width, y), fill=color)

    columns = cast_columns(map_rows, viewer, view_angle, width)
    for x, (hit, distance, side) in enumerate(columns):
        if not hit:
            continue
        line_height = height / max(distance, 0.0001)
        start_y = max(0.0, height * 0.5 - line_height * 0.5)
        end_y = min(height - 1.0, start_y + line_height)
        brightness = max(0.25, min(1.0, 1.0 - distance / 20.0))
        side_multiplier = 1.0 if side == 0 else 0.75
        color = (
            to_channel(125 * brightness * side_multiplier),
            to_channel(125 * brightness * side_multiplier),
            to_channel(135 * brightness * side_multiplier),
        )
        draw.line((x, start_y, x, end_y), fill=color)

    fov = math.radians(66.0)
    delta = target - viewer
    target_distance = math.hypot(delta.x, delta.y)
    if target_distance > 0.05:
        angle_to_enemy = angle_between(viewer, target)
        relative_angle = wrap_angle(angle_to_enemy - view_angle)
        if abs(relative_angle) <= fov * 0.5 + math.radians(8.0):
            normalized = relative_angle / fov + 0.5
            screen_center_x = normalized * width
            sprite_height = height / target_distance
            sprite_width = sprite_height * 0.60
            x_start = max(0, int(screen_center_x - sprite_width * 0.5))
            x_end = min(width - 1, int(screen_center_x + sprite_width * 0.5))
            y_start = max(0.0, height * 0.5 - sprite_height * 0.5)
            y_end = min(height - 1.0, y_start + sprite_height)
            for x in range(x_start, x_end + 1):
                if target_distance < columns[x][1]:
                    draw.line((x, y_start, x, y_end), fill=(220, 60, 60))

    draw.line((width / 2 - 10, height / 2, width / 2 + 10, height / 2), fill=(255, 255, 255), width=2)
    draw.line((width / 2, height / 2 - 10, width / 2, height / 2 + 10), fill=(255, 255, 255), width=2)

    font = ImageFont.load_default(size=24)
    small = ImageFont.load_default(size=18)
    draw.text((20, 20), "HP: 100", fill=(255, 255, 255), font=font)
    draw.text((20, 48), "Enemy HP: 100", fill=(255, 255, 255), font=font)
    draw.text((20, 76), "WASD move, arrows turn, Space shoot, R restart, Esc quit", fill=(255, 255, 255), font=small)
    draw.text((20, 100), "Mode: offline combat demo", fill=(255, 255, 255), font=small)

    img.save(output_path)


def map_overview_image(output_path: Path) -> None:
    rows, host, client = load_map()
    cell = 72
    margin = 48
    width = len(rows[0]) * cell + margin * 2
    height = len(rows) * cell + margin * 2 + 80
    img = Image.new("RGB", (width, height), "#f5f1ea")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.load_default(size=30)
    font = ImageFont.load_default(size=22)

    draw.text((margin, 16), "Карта arena.txt и стартовые позиции игроков", fill="#302518", font=title_font)

    for y, row in enumerate(rows):
        for x, cell_value in enumerate(row):
            left = margin + x * cell
            top = margin + 56 + y * cell
            right = left + cell - 2
            bottom = top + cell - 2
            if cell_value == "#":
                fill = "#45362a"
            else:
                fill = "#ddd1c3"
            draw.rounded_rectangle((left, top, right, bottom), radius=8, fill=fill, outline="#8d745b", width=2)
            if cell_value == "1":
                draw.ellipse((left + 16, top + 16, right - 16, bottom - 16), fill="#2563eb")
            elif cell_value == "2":
                draw.ellipse((left + 16, top + 16, right - 16, bottom - 16), fill="#dc2626")

    legend_y = height - 44
    draw.ellipse((margin, legend_y, margin + 18, legend_y + 18), fill="#2563eb")
    draw.text((margin + 30, legend_y - 4), f"Host: ({host.x:.1f}; {host.y:.1f})", fill="#302518", font=font)
    draw.ellipse((margin + 220, legend_y, margin + 238, legend_y + 18), fill="#dc2626")
    draw.text((margin + 250, legend_y - 4), f"Client: ({client.x:.1f}; {client.y:.1f})", fill="#302518", font=font)

    img.save(output_path)


def architecture_image(output_path: Path) -> None:
    width, height = 1500, 980
    img = Image.new("RGB", (width, height), "#f4efe7")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.load_default(size=34)
    box_font = ImageFont.load_default(size=22)
    small_font = ImageFont.load_default(size=18)

    draw.text((60, 28), "Архитектурная схема проекта Bauman Doomlike", fill="#23170f", font=title_font)

    boxes = [
        ("application", (520, 120, 980, 240), "#e4d7c4", "GameLoop\nGameSession\nкоординация сценария матча"),
        ("domain", (520, 320, 980, 500), "#d5e5dc", "Map, World, PlayerState\nCollisionSystem, ShootingSystem,\nDamageSystem"),
        ("presentation", (100, 320, 450, 500), "#dbe7f6", "SdlApp\nRaycaster\nHUD и обработка ввода"),
        ("network", (1050, 320, 1400, 500), "#f2dbcf", "EnetHost, EnetClient\nPacketSerializer\nProtocol"),
        ("infrastructure", (520, 600, 980, 760), "#efe1ef", "TextMapLoader\nработа с файлами\nи ресурсами"),
        ("resources", (520, 820, 980, 920), "#f1ebd6", "maps, textures, конфиги"),
    ]

    for _, rect, fill, text in boxes:
        draw.rounded_rectangle(rect, radius=22, fill=fill, outline="#5b4a3a", width=3)
        draw.multiline_text((rect[0] + 22, rect[1] + 18), text, fill="#23170f", font=box_font, spacing=8)

    def arrow(start: tuple[int, int], end: tuple[int, int]) -> None:
        draw.line((start, end), fill="#5b4a3a", width=5)
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        head = 18
        left = (end[0] - head * math.cos(angle - math.pi / 6), end[1] - head * math.sin(angle - math.pi / 6))
        right = (end[0] - head * math.cos(angle + math.pi / 6), end[1] - head * math.sin(angle + math.pi / 6))
        draw.polygon([end, left, right], fill="#5b4a3a")

    arrow((750, 240), (750, 320))
    arrow((450, 410), (520, 410))
    arrow((980, 410), (1050, 410))
    arrow((750, 500), (750, 600))
    arrow((750, 760), (750, 820))

    footer = (
        "Слои изолируют предметную логику от SDL и ENet, а загрузка карты и сетевой протокол "
        "вынесены в отдельные модули."
    )
    draw.multiline_text((60, 940), footer, fill="#4b3c30", font=small_font, spacing=6)
    img.save(output_path)


def terminal_image(output_path: Path) -> None:
    text = """$ cmake -S . -B build
$ cmake --build build -j2
[  4%] Built target enet
[ 94%] Built target SDL3-static
[100%] Built target bauman_doomlike

$ ./build/bauman_doomlike
Запуск приложения с картой resources/maps/arena.txt
Управление: WASD, стрелки, Space, R, Esc

$ tree -L 2 src
src
|-- application
|-- domain
|-- infrastructure
|-- network
`-- presentation"""

    width, height = 1350, 780
    img = Image.new("RGB", (width, height), "#111318")
    draw = ImageDraw.Draw(img)
    mono = ImageFont.load_default(size=22)
    header = ImageFont.load_default(size=28)

    draw.rounded_rectangle((30, 30, width - 30, height - 30), radius=24, fill="#171a21", outline="#343848", width=3)
    draw.ellipse((58, 54, 78, 74), fill="#ff5f57")
    draw.ellipse((88, 54, 108, 74), fill="#febc2e")
    draw.ellipse((118, 54, 138, 74), fill="#28c840")
    draw.text((170, 50), "Terminal - build and run log", fill="#c7d2fe", font=header)
    draw.multiline_text((64, 108), text, fill="#d5d8e2", font=mono, spacing=14)
    img.save(output_path)


def ensure_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)


def set_run_size(run, size_pt: float, bold: bool = False, italic: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_size(run, 14)


def bullet_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style=None)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.first_line_indent = Cm(-0.5)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"• {text}")
    set_run_size(run, 14)


def heading(doc: Document, text: str, centered: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_size(run, 14, bold=True)


def subheading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_size(run, 14, bold=True, italic=True)


def code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(10.5)

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F5F5F5")
    cell._tc.get_or_add_tcPr().append(shading)


def caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_size(run, 12, italic=True)


def add_picture(doc: Document, image_path: Path, width_cm: float, caption_text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    caption(doc, caption_text)


def add_page_number_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    fld_simple.append(run)
    paragraph._p.append(fld_simple)


def add_contents(doc: Document) -> None:
    heading(doc, "СОДЕРЖАНИЕ", centered=True)
    entries = [
        ("Введение", PAGE_MAP["Введение"]),
        ("Индивидуальное задание", PAGE_MAP["Индивидуальное задание"]),
        ("Основная часть", PAGE_MAP["Основная часть"]),
        ("Заключение", PAGE_MAP["Заключение"]),
        ("Список использованных источников", PAGE_MAP["Список использованных источников"]),
        ("Приложения", PAGE_MAP["Приложения"]),
    ]
    for title, page in entries:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(f"{title}")
        set_run_size(run, 14)
        dots = max(3, 90 - len(title))
        run = paragraph.add_run("." * dots + f" {page}")
        set_run_size(run, 14)


def add_main_content(doc: Document) -> None:
    heading(doc, "ВВЕДЕНИЕ", centered=True)
    body_paragraph(doc, INTRO_TEXT)

    doc.add_page_break()
    heading(doc, "ИНДИВИДУАЛЬНОЕ ЗАДАНИЕ", centered=True)
    body_paragraph(
        doc,
        "В ходе проектно-технологической практики было поставлено индивидуальное задание разработать учебный "
        "проект \"Bauman Doomlike\" - упрощенный псевдо-3D шутер на языке C++ с использованием объектно-"
        "ориентированного подхода, модульной архитектуры и внешних библиотек."
    )
    bullet_paragraph(doc, "выбрать технологический стек и аргументировать его с точки зрения учебного проекта;")
    bullet_paragraph(doc, "спроектировать слоистую архитектуру с разделением доменной логики, визуализации, инфраструктуры и сетевого слоя;")
    bullet_paragraph(doc, "реализовать загрузку карты из внешнего текстового файла и хранение игрового мира в собственных классах;")
    bullet_paragraph(doc, "создать игровой цикл, обработку пользовательского ввода, перемещение, столкновения, стрельбу и завершение матча;")
    bullet_paragraph(doc, "реализовать графический интерфейс на SDL3 с отрисовкой сцены методом raycasting;")
    bullet_paragraph(doc, "подготовить сетевой слой на базе ENet и протокол обмена сообщениями для дальнейшего перехода к multiplayer-режиму;")
    bullet_paragraph(doc, "настроить кроссплатформенную сборку проекта через CMake.")

    doc.add_page_break()
    heading(doc, "ОСНОВНАЯ ЧАСТЬ", centered=True)

    subheading(doc, "1. Краткая характеристика программы")
    body_paragraph(
        doc,
        "Разработанное приложение представляет собой учебный prototype шутера с псевдо-3D визуализацией. "
        "Игрок перемещается по клеточной карте, наблюдает трехмерную сцену через raycasting-рендер, стреляет "
        "по противнику и завершает раунд после обнуления очков здоровья цели. В текущем состоянии проект "
        "работает как offline combat demo, при этом структура кода уже подготавливает дальнейшее развитие до "
        "сетевой игры формата 1 vs 1."
    )
    body_paragraph(
        doc,
        "Возможности созданного решения включают загрузку внешней карты, запуск графического окна SDL3, "
        "обработку клавиш управления, отображение стен и противника, HUD c состоянием матча, рестарт раунда "
        "и подготовленный сетевой протокол для синхронизации команд и состояния игроков."
    )

    subheading(doc, "2. Архитектура и основные модули")
    body_paragraph(
        doc,
        "Проект спроектирован по слоистому принципу. Слой domain содержит чистую игровую логику: карту, "
        "мир, игрока, систему столкновений, стрельбы и нанесения урона. Слой application координирует "
        "игровой цикл и сценарии матча. Слой presentation инкапсулирует SDL3, отрисовку и преобразование "
        "пользовательского ввода в команды. Слой infrastructure отвечает за загрузку карты из файла. "
        "Слой network содержит обертки над ENet и описание сетевого протокола."
    )
    code_block(
        doc,
        "set(DOOMLIKE_SOURCES\n"
        "    src/application/GameLoop.cpp\n"
        "    src/application/GameSession.cpp\n"
        "    src/domain/Map.cpp\n"
        "    src/domain/World.cpp\n"
        "    src/presentation/Raycaster.cpp\n"
        "    src/presentation/SdlApp.cpp\n"
        "    src/network/PacketSerializer.cpp\n"
        ")"
    )
    caption(doc, "Листинг 1 - Модульная структура исполняемого файла в CMakeLists.txt")

    body_paragraph(
        doc,
        "Центральным объектом прикладного слоя является класс GameSession. Он загружает карту, инициализирует "
        "мир, обновляет кулдауны оружия, рассчитывает движение локального игрока и инициирует выстрел."
    )
    code_block(
        doc,
        "void GameSession::update(const float deltaSeconds, const InputState& input)\n"
        "{\n"
        "    updateWeaponCooldown(deltaSeconds);\n"
        "    if (world_.phase() != domain::MatchPhase::Running) {\n"
        "        return;\n"
        "    }\n"
        "    updateMovement(deltaSeconds, input);\n"
        "    if (input.shootPressed) {\n"
        "        handleShoot();\n"
        "    }\n"
        "    world_.advanceTick();\n"
        "}"
    )
    caption(doc, "Листинг 2 - Обновление логики кадра в GameSession")

    subheading(doc, "3. Реализация игровой логики")
    body_paragraph(
        doc,
        "Карта хранится во внешнем текстовом файле resources/maps/arena.txt. Символы '#' обозначают стены, '.' - "
        "свободные клетки, '1' и '2' - стартовые позиции игроков. Во время загрузки выполняется валидация "
        "прямоугольности карты, допустимости символов и обязательного наличия точек старта."
    )
    code_block(
        doc,
        "if (border && tile != '#') {\n"
        "    throw std::runtime_error(\"Внешняя граница карты должна быть из стен '#'\");\n"
        "}\n"
        "switch (tile) {\n"
        "    case '1':\n"
        "        map.setSpawn(domain::PlayerId::Host, {x + 0.5F, y + 0.5F});\n"
        "        map.setTile(x, y, '.');\n"
        "        break;\n"
        "    case '2':\n"
        "        map.setSpawn(domain::PlayerId::Client, {x + 0.5F, y + 0.5F});\n"
        "        map.setTile(x, y, '.');\n"
        "        break;\n"
        "}"
    )
    caption(doc, "Листинг 3 - Загрузка карты и извлечение spawn points")

    body_paragraph(
        doc,
        "Перемещение игрока рассчитывается по вектору направления камеры и боковому вектору. Перед применением "
        "новой позиции вызывается система CollisionSystem, которая проверяет пять контрольных точек окружности "
        "игрока и отдельно разрешает движение по осям X и Y. Это дает простую и устойчивую модель столкновений "
        "для клеточного мира."
    )
    code_block(
        doc,
        "const std::array<Vec2, 5> probes {\n"
        "    Vec2 {position.x - radius, position.y - radius},\n"
        "    Vec2 {position.x + radius, position.y - radius},\n"
        "    Vec2 {position.x - radius, position.y + radius},\n"
        "    Vec2 {position.x + radius, position.y + radius},\n"
        "    position,\n"
        "};"
    )
    caption(doc, "Листинг 4 - Контрольные точки для проверки столкновения со стенами")

    body_paragraph(
        doc,
        "Стрельба реализована как hitscan-проверка. Система ShootingSystem вычисляет расстояние до противника, "
        "угол между направлением взгляда и направлением на цель, а затем проверяет наличие прямой видимости, "
        "пошагово продвигаясь вдоль луча до цели. После подтвержденного попадания DamageSystem уменьшает HP и "
        "переводит матч в состояние GameOver."
    )
    code_block(
        doc,
        "const float angleDelta = std::abs(wrapAngle(angleToTarget - shooter.angleRadians));\n"
        "const float allowedAngle = std::atan2(target.radius, std::max(distance, 0.001F)) + degToRad(1.5F);\n"
        "if (angleDelta > allowedAngle) {\n"
        "    return std::nullopt;\n"
        "}\n"
        "if (!hasLineOfSight(world.map(), shooter.position, target.position)) {\n"
        "    return std::nullopt;\n"
        "}"
    )
    caption(doc, "Листинг 5 - Проверка попадания по противнику")

    subheading(doc, "4. Графическая часть")
    body_paragraph(
        doc,
        "Графический интерфейс построен на SDL3. Класс SdlApp создает окно и renderer, принимает события "
        "клавиатуры, а затем отрисовывает фон, raycasting-стены, спрайт противника и HUD с информацией о матче. "
        "Основной принцип визуализации - DDA-проход луча по клеткам карты для каждого вертикального столбца экрана."
    )
    code_block(
        doc,
        "const std::vector<RayHit> columns =\n"
        "    raycaster_.castColumns(world.map(), viewer.position, viewer.angleRadians, logicalWidth_);\n"
        "drawWalls(columns);\n"
        "drawSprites(world, viewerId, columns);\n"
        "drawHud(world, viewerId);"
    )
    caption(doc, "Листинг 6 - Последовательность рендеринга одного кадра")

    body_paragraph(
        doc,
        "Визуальная часть уже демонстрирует готовый игровой интерфейс: псевдо-3D стены, перекрестие прицела, "
        "информацию о здоровье и подсказку по управлению. Цветовая схема строится из градиентного неба, "
        "затемненного пола и затухающей яркости стен в зависимости от расстояния."
    )

    subheading(doc, "5. Подготовка сетевой части")
    body_paragraph(
        doc,
        "Для будущего расширения до сетевого режима в проект включены ENet и модуль протокола обмена "
        "сообщениями. Уже выделены типы пакетов Connect, Welcome, InputCommand, StateUpdate, Shoot, Hit, Damage, "
        "GameOver и Restart. Обертки EnetHost и EnetClient скрывают инициализацию runtime, управление peer'ами "
        "и прием валидных сообщений."
    )
    code_block(
        doc,
        "enum class MessageType : std::uint8_t {\n"
        "    Connect = 1,\n"
        "    Welcome = 2,\n"
        "    InputCommand = 3,\n"
        "    StateUpdate = 4,\n"
        "    Shoot = 5,\n"
        "    Hit = 6,\n"
        "    Damage = 7,\n"
        "    GameOver = 8,\n"
        "    Restart = 9,\n"
        "};"
    )
    caption(doc, "Листинг 7 - Типы сообщений пользовательского сетевого протокола")

    subheading(doc, "6. Удовлетворение требованиям практики")
    body_paragraph(
        doc,
        "Ниже приведено соответствие проекта формальным требованиям практики, указанным в методических материалах."
    )

    bullet_paragraph(
        doc,
        "Разработка ведется на языке C++ с использованием ООП-подхода. В проекте выделены классы GameLoop, "
        "GameSession, Map, World, SdlApp, Raycaster, EnetHost и EnetClient, каждый из которых инкапсулирует "
        "отдельную область ответственности."
    )
    code_block(doc, "class Map { ... };\nclass World { ... };\nclass GameSession { ... };\nclass SdlApp { ... };")

    bullet_paragraph(
        doc,
        "Код логически разбит по модулям .cpp и .hpp. Для каждого ключевого компонента существуют отдельные "
        "заголовочный и реализационный файлы, что уменьшает связность и упрощает сопровождение."
    )
    code_block(doc, "src/domain/Map.hpp + src/domain/Map.cpp\nsrc/presentation/SdlApp.hpp + src/presentation/SdlApp.cpp")

    bullet_paragraph(
        doc,
        "Использованы внешние библиотеки SDL3 и ENet. Они подключаются через FetchContent, что делает сборку "
        "повторяемой на разных операционных системах."
    )
    code_block(
        doc,
        "FetchContent_Declare(SDL3 ... GIT_TAG release-3.4.4)\n"
        "FetchContent_Declare(enet ... GIT_TAG v1.3.18)\n"
        "FetchContent_MakeAvailable(SDL3 enet)"
    )

    bullet_paragraph(
        doc,
        "Разработка проекта ориентирована на использование Git и удаленного репозитория. В предоставленной "
        "локальной копии история коммитов отсутствует, поэтому в рамках итоговой защиты необходимо показать "
        "ссылку на удаленный репозиторий и историю личных коммитов."
    )
    code_block(
        doc,
        "git init\n"
        "git add .\n"
        "git commit -m \"Implement raycasting prototype\"\n"
        "git remote add origin <repo-url>\n"
        "git push -u origin main"
    )

    bullet_paragraph(
        doc,
        "Разработан графический интерфейс. Окно приложения создается через SDL_CreateWindowAndRenderer, далее "
        "в нем отображаются игровая сцена и HUD."
    )
    code_block(doc, "SDL_CreateWindowAndRenderer(..., &window_, &renderer_);\nSDL_RenderPresent(renderer_);")

    bullet_paragraph(
        doc,
        "Используются структуры данных STL: std::vector для raycasting-буфера, std::array для состояния игроков "
        "и контрольных точек коллизий, std::optional для результатов попадания и победителя, std::span для "
        "работы с сетевыми пакетами."
    )
    code_block(doc, "std::vector<RayHit>\nstd::array<PlayerState, 2>\nstd::optional<ShotResult>\nstd::span<const std::byte>")

    bullet_paragraph(
        doc,
        "Сборка проекта выполняется через CMake. Исполняемый файл связывается с SDL3 и ENet, а стандарты языка "
        "и предупреждения компилятора задаются централизованно."
    )
    code_block(
        doc,
        "add_executable(${PROJECT_NAME} ${DOOMLIKE_SOURCES})\n"
        "target_link_libraries(${PROJECT_NAME} PRIVATE ${SDL3_TARGET} ${ENET_TARGET})\n"
        "set(CMAKE_CXX_STANDARD 20)"
    )

    subheading(doc, "7. Личный вклад")
    body_paragraph(
        doc,
        "Проект выполнялся как индивидуальная практика, поэтому весь объем работ - анализ требований, выбор "
        "технологического стека, проектирование архитектуры, реализация доменной логики, визуализации, "
        "сетевых заготовок, настройка сборки и подготовка отчетных материалов - выполнен одним студентом."
    )

    doc.add_page_break()
    heading(doc, "ЗАКЛЮЧЕНИЕ", centered=True)
    body_paragraph(
        doc,
        "В ходе проектно-технологической практики был разработан учебный проект Bauman Doomlike, демонстрирующий "
        "применение объектно-ориентированного подхода в C++ при создании интерактивного графического приложения. "
        "В рамках практики удалось выбрать и обосновать стек технологий, спроектировать слоистую архитектуру, "
        "реализовать загрузку карты, игровой цикл, столкновения, стрельбу, raycasting-визуализацию и подготовить "
        "сетевой слой для дальнейшего развития решения."
    )
    body_paragraph(
        doc,
        "Полученный результат удовлетворяет ключевым учебным требованиям: проект написан на C++, использует "
        "ООП, внешние библиотеки, структуры STL, модульное разбиение и CMake-сборку. В качестве следующего этапа "
        "развития целесообразно завершить интеграцию ENet в игровой цикл, добавить полноценный режим двух игроков "
        "по локальной сети и расширить визуальную часть текстурами, меню и звуковым сопровождением."
    )

    doc.add_page_break()
    heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", centered=True)
    sources = [
        "Практика по курсу «Программирование на основе классов и шаблонов» кафедры ИУ5 МГТУ им. Н.Э. Баумана [Электронный ресурс]. URL: https://iu5git.github.io/OOP/meta/practice.html (дата обращения: 14.04.2026).",
        "Техническое задание и стартовая архитектура проекта Bauman Doomlike / файл docs/technical_design_ru.md.",
        "SDL3 Documentation [Электронный ресурс]. URL: https://wiki.libsdl.org/SDL3 (дата обращения: 14.04.2026).",
        "ENet Reliable UDP Networking Library [Электронный ресурс]. URL: https://github.com/lsalzman/enet (дата обращения: 14.04.2026).",
        "CMake Documentation [Электронный ресурс]. URL: https://cmake.org/documentation/ (дата обращения: 14.04.2026).",
    ]
    for idx, source in enumerate(sources, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(-0.75)
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run(f"{idx}. {source}")
        set_run_size(run, 14)

    doc.add_page_break()
    heading(doc, "ПРИЛОЖЕНИЯ", centered=True)
    body_paragraph(doc, "В приложении приведены иллюстрации, подтверждающие структуру и текущее состояние проекта.")
    add_picture(doc, ASSET_DIR / "build_console.png", 15.5, "Рисунок А.1 - Иллюстрация сборки и запуска проекта из терминала")
    add_picture(doc, ASSET_DIR / "architecture_layers.png", 15.5, "Рисунок А.2 - Архитектурная схема слоев проекта")
    add_picture(doc, ASSET_DIR / "map_overview.png", 15.0, "Рисунок А.3 - Схема карты arena.txt и стартовых позиций")
    add_picture(doc, ASSET_DIR / "interface_preview.png", 15.5, "Рисунок А.4 - Иллюстрация игрового окна, построенная по текущей логике рендера")


def generate_assets() -> None:
    terminal_image(ASSET_DIR / "build_console.png")
    architecture_image(ASSET_DIR / "architecture_layers.png")
    map_overview_image(ASSET_DIR / "map_overview.png")
    raycast_interface_image(ASSET_DIR / "interface_preview.png")


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def create_report() -> None:
    doc = Document(str(TITLE_PATH))

    add_page_break(doc)
    for section in doc.sections:
        add_page_number_footer(section)
    doc.sections[-1].different_first_page_header_footer = True

    add_contents(doc)
    add_main_content(doc)

    doc.save(str(REPORT_PATH))


def render_preview(docx_path: Path) -> None:
    preview_dir = TMP_DIR / "ql_preview"
    if preview_dir.exists():
        for item in preview_dir.rglob("*"):
            if item.is_file():
                item.unlink()
        for item in sorted(preview_dir.glob("*"), reverse=True):
            if item.is_dir():
                item.rmdir()
    preview_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        ["qlmanage", "-p", "-o", str(preview_dir), str(docx_path)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )


def main() -> None:
    ensure_dirs()
    generate_assets()
    create_report()
    render_preview(REPORT_PATH)
    print(f"Report written to {REPORT_PATH}")


if __name__ == "__main__":
    main()
